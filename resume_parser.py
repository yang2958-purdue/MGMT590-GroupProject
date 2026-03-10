"""
Resume parser module for handling resume input.
Supports file upload and pasted text with validation.
Handles TXT, PDF, and DOCX file formats.
"""

from pathlib import Path
from typing import Optional
from utils import normalize_text, sanitize_input, validate_resume_text

# Import PDF and DOCX parsers
try:
    from PyPDF2 import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class ResumeParser:
    """
    Handles resume input from file or direct text entry.
    Provides validation and normalization.
    Supports TXT, PDF, and DOCX file formats.
    """
    
    def __init__(self):
        self.resume_text: str = ""
        self.source: str = ""
    
    def _extract_text_from_pdf(self, file_path: Path) -> tuple[bool, str, str]:
        """
        Extract text from a PDF file.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Tuple of (success, content, error_message)
        """
        if not PDF_AVAILABLE:
            return False, "", "PDF support not available. Please install PyPDF2: pip install PyPDF2"
        
        try:
            reader = PdfReader(str(file_path))
            
            # Extract text from all pages
            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            
            content = "\n".join(text_parts)
            
            if not content.strip():
                return False, "", "PDF appears to be empty or contains no extractable text."
            
            return True, content, ""
            
        except Exception as e:
            return False, "", f"Error reading PDF: {str(e)}"
    
    def _extract_text_from_docx(self, file_path: Path) -> tuple[bool, str, str]:
        """
        Extract text from a Word document (.docx).
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Tuple of (success, content, error_message)
        """
        if not DOCX_AVAILABLE:
            return False, "", "DOCX support not available. Please install python-docx: pip install python-docx"
        
        try:
            doc = Document(str(file_path))
            
            # Extract text from all paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            content = "\n".join(text_parts)
            
            if not content.strip():
                return False, "", "Word document appears to be empty."
            
            return True, content, ""
            
        except Exception as e:
            return False, "", f"Error reading Word document: {str(e)}"
    
    def _extract_text_from_txt(self, file_path: Path) -> tuple[bool, str, str]:
        """
        Extract text from a text file.
        
        Args:
            file_path: Path to the TXT file
            
        Returns:
            Tuple of (success, content, error_message)
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            if not content.strip():
                return False, "", "Text file is empty."
            
            return True, content, ""
            
        except UnicodeDecodeError:
            return False, "", "Unable to read file. Please ensure it's UTF-8 encoded."
        except Exception as e:
            return False, "", f"Error reading text file: {str(e)}"
    
    def load_from_file(self, file_path: str) -> tuple[bool, str]:
        """
        Load resume from a file (TXT, PDF, or DOCX).
        
        Args:
            file_path: Path to the resume file
            
        Returns:
            Tuple of (success, message)
        """
        try:
            path = Path(file_path)
            
            # Check if file exists
            if not path.exists():
                return False, f"File not found: {file_path}"
            
            # Check if it's a file
            if not path.is_file():
                return False, f"Path is not a file: {file_path}"
            
            # Determine file type and extract text
            file_extension = path.suffix.lower()
            
            if file_extension == '.pdf':
                success, content, error_msg = self._extract_text_from_pdf(path)
            elif file_extension in ['.docx', '.doc']:
                if file_extension == '.doc':
                    return False, "Old .doc format not supported. Please save as .docx"
                success, content, error_msg = self._extract_text_from_docx(path)
            elif file_extension == '.txt':
                success, content, error_msg = self._extract_text_from_txt(path)
            else:
                return False, f"Unsupported file type: {file_extension}. Please use .txt, .pdf, or .docx"
            
            # Check if extraction was successful
            if not success:
                return False, error_msg
            
            # Sanitize and normalize
            content = sanitize_input(content)
            
            # Validate
            is_valid, error_msg = validate_resume_text(content)
            if not is_valid:
                return False, error_msg
            
            # Store resume
            self.resume_text = content
            self.source = f"file: {path.name} ({file_extension[1:].upper()})"
            
            return True, f"Successfully loaded resume from {path.name}"
            
        except PermissionError:
            return False, "Permission denied. Unable to read the file."
        except Exception as e:
            return False, f"Error reading file: {str(e)}"
    
    def load_from_text(self, text: str) -> tuple[bool, str]:
        """
        Load resume from pasted text.
        
        Args:
            text: Resume text content
            
        Returns:
            Tuple of (success, message)
        """
        try:
            # Check if text is empty
            if not text or not text.strip():
                return False, "Resume text is empty. Please provide your resume content."
            
            # Sanitize and normalize
            content = sanitize_input(text)
            
            # Validate
            is_valid, error_msg = validate_resume_text(content)
            if not is_valid:
                return False, error_msg
            
            # Store resume
            self.resume_text = content
            self.source = "pasted text"
            
            return True, "Successfully loaded resume from pasted text"
            
        except Exception as e:
            return False, f"Error processing text: {str(e)}"
    
    def get_resume_text(self) -> str:
        """
        Get the loaded resume text.
        
        Returns:
            Resume text or empty string if not loaded
        """
        return self.resume_text
    
    def get_resume_preview(self, max_chars: int = 200) -> str:
        """
        Get a preview of the loaded resume.
        
        Args:
            max_chars: Maximum characters to return
            
        Returns:
            Preview text
        """
        if not self.resume_text:
            return "No resume loaded"
        
        preview = self.resume_text[:max_chars]
        if len(self.resume_text) > max_chars:
            preview += "..."
        
        return preview
    
    def is_loaded(self) -> bool:
        """
        Check if a resume has been loaded.
        
        Returns:
            True if resume is loaded, False otherwise
        """
        return bool(self.resume_text)
    
    def clear(self) -> None:
        """
        Clear the loaded resume.
        """
        self.resume_text = ""
        self.source = ""
    
    def get_word_count(self) -> int:
        """
        Get word count of loaded resume.
        
        Returns:
            Number of words
        """
        if not self.resume_text:
            return 0
        return len(self.resume_text.split())
    
    def get_char_count(self) -> int:
        """
        Get character count of loaded resume.
        
        Returns:
            Number of characters
        """
        return len(self.resume_text)

