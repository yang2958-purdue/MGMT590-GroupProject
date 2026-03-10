"""
Resume tailoring service - Customizes resumes for specific job descriptions using Claude AI
"""
import os
import re
from typing import Dict, List
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import anthropic
import hashlib
from datetime import datetime


class ResumeTailor:
    """Tailors resumes to match specific job descriptions"""
    
    def __init__(self):
        self.anthropic_client = None
        anthropic_key = os.getenv("ANTHROPIC_API_KEY")
        if anthropic_key:
            self.anthropic_client = anthropic.Anthropic(api_key=anthropic_key)
        
        self.output_dir = "tailored_resumes"
        os.makedirs(self.output_dir, exist_ok=True)
    
    def tailor_resume(
        self,
        resume_text: str,
        job_description: str,
        output_format: str = "docx"
    ) -> Dict:
        """
        Tailor a resume for a specific job
        
        Args:
            resume_text: Original resume text
            job_description: Target job description
            output_format: Output format (docx, pdf, or txt)
            
        Returns:
            Dictionary with tailored resume and metadata
        """
        if not self.anthropic_client:
            raise ValueError("Claude API key not configured. Cannot tailor resume.")
        
        # Generate tailored content using Claude
        tailored_text, changes = self._generate_tailored_content(
            resume_text, job_description
        )
        
        # Generate unique ID for tailored resume
        tailored_id = self._generate_id(tailored_text)
        
        # Create output file
        filename = f"tailored_resume_{tailored_id}.{output_format}"
        filepath = os.path.join(self.output_dir, filename)
        
        if output_format == "docx":
            self._create_docx(tailored_text, filepath)
        elif output_format == "txt":
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(tailored_text)
        elif output_format == "pdf":
            # PDF generation would require additional library (e.g., reportlab)
            # For now, create DOCX and mention PDF conversion
            docx_path = filepath.replace('.pdf', '.docx')
            self._create_docx(tailored_text, docx_path)
            filepath = docx_path
        
        return {
            "tailored_resume_id": tailored_id,
            "tailored_text": tailored_text,
            "changes_made": changes,
            "download_url": f"/api/resume/download/{tailored_id}",
            "filepath": filepath,
            "format": output_format
        }
    
    def _generate_tailored_content(
        self,
        resume_text: str,
        job_description: str
    ) -> tuple[str, List[str]]:
        """
        Use Claude to generate tailored resume content
        """
        prompt = f"""You are an expert resume writer. Tailor the following resume to better match the job description.

ORIGINAL RESUME:
{resume_text}

TARGET JOB DESCRIPTION:
{job_description[:2000]}

Instructions:
1. Rewrite bullet points to emphasize relevant experience and skills
2. Adjust the professional summary to align with the role
3. Reorder or emphasize skills that match the job requirements
4. Keep all information truthful - do not add fake experience
5. Maintain the overall structure and formatting cues (like section headers)
6. Keep it concise and impactful

Provide two things:
1. The complete tailored resume text
2. A bulleted list of key changes made (3-5 items)

Format your response EXACTLY as:
TAILORED_RESUME:
[full tailored resume text here]

CHANGES_MADE:
- Change 1
- Change 2
- Change 3"""

        try:
            message = self.anthropic_client.messages.create(
                model="claude-3-5-sonnet-20241022",
                max_tokens=4096,
                messages=[{"role": "user", "content": prompt}]
            )
            
            response_text = message.content[0].text.strip()
            
            # Parse response
            tailored_resume, changes = self._parse_tailoring_response(response_text)
            
            return tailored_resume, changes
        
        except Exception as e:
            print(f"Claude AI tailoring failed: {str(e)}")
            # Return original resume with error message
            return resume_text, [f"Tailoring failed: {str(e)}"]
    
    def _parse_tailoring_response(self, response: str) -> tuple[str, List[str]]:
        """Parse Claude's response into resume and changes"""
        # Split by markers
        parts = response.split("CHANGES_MADE:")
        
        if len(parts) == 2:
            resume_part = parts[0].replace("TAILORED_RESUME:", "").strip()
            changes_part = parts[1].strip()
            
            # Parse changes (bulleted list)
            changes = []
            for line in changes_part.split('\n'):
                line = line.strip()
                if line.startswith('-') or line.startswith('•'):
                    changes.append(line[1:].strip())
                elif line and not line.startswith('TAILORED_RESUME'):
                    changes.append(line)
            
            return resume_part, changes
        else:
            # Fallback if parsing fails
            return response, ["Resume tailored by AI"]
    
    def _create_docx(self, text: str, filepath: str):
        """
        Create a formatted DOCX file from text
        Attempts to preserve structure from the text
        """
        doc = Document()
        
        # Set document margins (narrow)
        sections = doc.sections
        for section in sections:
            section.top_margin = Pt(36)
            section.bottom_margin = Pt(36)
            section.left_margin = Pt(72)
            section.right_margin = Pt(72)
        
        # Parse text into sections
        lines = text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Detect headers (ALL CAPS or Title Case with no punctuation at end)
            if self._is_header(line):
                # Add header
                p = doc.add_paragraph(line)
                p.style = 'Heading 1'
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.runs[0]
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
            
            # Detect bullet points
            elif line.startswith('-') or line.startswith('•') or line.startswith('*'):
                p = doc.add_paragraph(line[1:].strip(), style='List Bullet')
                run = p.runs[0]
                run.font.size = Pt(11)
            
            # Regular text
            else:
                p = doc.add_paragraph(line)
                run = p.runs[0]
                run.font.size = Pt(11)
        
        doc.save(filepath)
    
    def _is_header(self, line: str) -> bool:
        """Detect if a line is likely a section header"""
        # Check if line is all uppercase
        if line.isupper() and len(line) > 2:
            return True
        
        # Check if line matches common header patterns
        header_keywords = [
            'EXPERIENCE', 'EDUCATION', 'SKILLS', 'SUMMARY', 'OBJECTIVE',
            'PROJECTS', 'CERTIFICATIONS', 'PUBLICATIONS', 'AWARDS',
            'Professional Experience', 'Work Experience', 'Technical Skills'
        ]
        
        for keyword in header_keywords:
            if keyword.lower() in line.lower() and len(line) < 50:
                return True
        
        return False
    
    def _generate_id(self, text: str) -> str:
        """Generate unique ID from text"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        text_hash = hashlib.md5(text.encode()).hexdigest()[:8]
        return f"{timestamp}_{text_hash}"
    
    def get_tailored_resume(self, tailored_id: str) -> str:
        """Retrieve a tailored resume by ID"""
        # Look for file with matching ID
        for filename in os.listdir(self.output_dir):
            if tailored_id in filename:
                filepath = os.path.join(self.output_dir, filename)
                return filepath
        
        raise FileNotFoundError(f"Tailored resume {tailored_id} not found")


# Singleton instance
resume_tailor = ResumeTailor()

